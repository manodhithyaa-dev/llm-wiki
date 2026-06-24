import os
import re
from pathlib import Path
from datetime import datetime

import fitz
from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from pipelines.utils import ensure_dir, append_log
from pipelines.ingest_image import ingest_image

DOCUMENTS_DIR = "raw/documents"

IMAGE_EXT_MAP = {
    "jpg": "jpg", "jpeg": "jpg", "png": "png", "gif": "gif",
    "bmp": "bmp", "tif": "tiff", "tiff": "tiff", "webp": "webp",
}


def _has_list_numbering(para) -> bool:
    num_pr = para._element.find(qn("w:pPr") + "/" + qn("w:numPr"))
    return num_pr is not None


def _get_list_level(para) -> int:
    num_pr = para._element.find(qn("w:pPr") + "/" + qn("w:numPr"))
    if num_pr is not None:
        ilvl = num_pr.find(qn("w:ilvl"))
        if ilvl is not None and ilvl.get(qn("w:val")) is not None:
            return int(ilvl.get(qn("w:val")))
    return 0


def _is_bullet(para) -> bool:
    num_pr = para._element.find(qn("w:pPr") + "/" + qn("w:numPr"))
    if num_pr is not None:
        num_id = num_pr.find(qn("w:numId"))
        if num_id is not None and num_id.get(qn("w:val")) is not None:
            return True
    style_name = para.style.name if para.style else ""
    return "List Bullet" in style_name


def _extract_pdf_images(doc, pdf_name: str, page_num: int) -> list:
    images_dir = Path(f"processed/documents/{pdf_name}/images")
    ensure_dir(str(images_dir))

    page = doc[page_num]
    image_list = page.get_images(full=True)
    extracted = []

    for img_idx, img in enumerate(image_list):
        xref = img[0]
        try:
            base = doc.extract_image(xref)
            ext = base["ext"]
            ext = IMAGE_EXT_MAP.get(ext, ext)
            img_file = images_dir / f"p{page_num+1:03d}_img{img_idx+1:03d}.{ext}"
            with open(img_file, "wb") as f:
                f.write(base["image"])
            extracted.append(img_file)
        except Exception:
            continue

    return extracted


def ingest_pdf(pdf_path: str) -> None:
    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
    out_dir = f"processed/documents/{pdf_name}/pages"
    ensure_dir(out_dir)

    doc = fitz.open(pdf_path)
    total_pages = len(doc)

    for page_num in range(total_pages):
        page = doc[page_num]
        text = page.get_text()
        page_file = f"{out_dir}/p{page_num+1:03d}.md"

        images = _extract_pdf_images(doc, pdf_name, page_num)

        image_refs = ""
        for img_file in images:
            rel_path = f"../images/{img_file.name}"
            image_refs += f"\n![Page {page_num+1} Image]({rel_path})\n"
            try:
                ingest_image(str(img_file))
            except Exception:
                pass

        with open(page_file, "w", encoding="utf-8") as f:
            f.write(f"# Page {page_num+1}\n\n")
            f.write(text)
            if image_refs:
                f.write(image_refs)

    doc.close()
    print(f"[PDF] Processed {pdf_name} ({total_pages} pages, extracted images)")
    append_log(f"Ingested PDF: {pdf_name} ({total_pages} pages)")


def _generate_source_metadata(source_path: str, doc_name: str) -> None:
    source_path_obj = Path(source_path)
    stat = source_path_obj.stat()

    metadata = f"""# Source Metadata

**File:** {source_path_obj.name}
**Type:** {source_path_obj.suffix.upper().lstrip('.')}
**Size:** {stat.st_size:,} bytes
**Created:** {datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M:%S")}
**Modified:** {datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")}
**Processed:** {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
"""

    meta_path = f"processed/documents/{doc_name}/source_metadata.md"
    with open(meta_path, "w", encoding="utf-8") as f:
        f.write(metadata)


def _table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    lines = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    for row in rows[1:]:
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _extract_docx_images(docx_path: str, doc_name: str) -> list:
    import zipfile

    images_dir = Path(f"processed/documents/{doc_name}/images")
    ensure_dir(str(images_dir))

    extracted = []
    count = 0

    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            for name in z.namelist():
                if name.startswith("word/media/") or name.startswith("media/"):
                    data = z.read(name)
                    ext = Path(name).suffix.lstrip(".").lower()
                    ext = IMAGE_EXT_MAP.get(ext, ext)
                    count += 1
                    img_file = images_dir / f"docx_img_{count:03d}.{ext}"
                    with open(img_file, "wb") as f:
                        f.write(data)
                    extracted.append(img_file)
                    try:
                        ingest_image(str(img_file))
                    except Exception:
                        pass
    except Exception:
        pass

    return extracted


def ingest_docx(docx_path: str) -> None:
    doc_name = os.path.splitext(os.path.basename(docx_path))[0]
    out_dir = f"processed/documents/{doc_name}/pages"
    ensure_dir(out_dir)

    extracted_images = _extract_docx_images(docx_path, doc_name)

    document = Document(docx_path)
    page_file = f"{out_dir}/p001.md"

    lines = []
    lines.append(f"# {doc_name}\n")

    if extracted_images:
        lines.append("## Embedded Images\n")
        for img_file in extracted_images:
            rel_path = f"../images/{img_file.name}"
            lines.append(f"![{img_file.name}]({rel_path})")
        lines.append("")

    body = document.element.body
    para_index = 0
    table_index = 0
    total_children = len(list(body))

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            if para_index < len(document.paragraphs):
                para = document.paragraphs[para_index]
                para_index += 1
                text = para.text.strip()
                style_name = para.style.name if para.style else "Normal"

                if not text:
                    continue

                if style_name.startswith("Heading"):
                    try:
                        level_str = style_name.replace("Heading", "").strip()
                        level = int(level_str) if level_str else 1
                    except ValueError:
                        level = 1
                    level = max(1, min(level, 6))
                    lines.append(f"{'#' * level} {text}")
                elif _has_list_numbering(para):
                    level = _get_list_level(para)
                    indent = "  " * level
                    prefix = "-" if _is_bullet(para) else "1."
                    lines.append(f"{indent}{prefix} {text}")
                else:
                    lines.append(text)
                    lines.append("")

        elif tag == "tbl":
            if table_index < len(document.tables):
                table = document.tables[table_index]
                table_index += 1
                md_table = _table_to_markdown(table)
                if md_table:
                    lines.append("")
                    lines.append(md_table)
                    lines.append("")

    # Remove trailing blank lines
    while lines and lines[-1] == "":
        lines.pop()

    with open(page_file, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    _generate_source_metadata(docx_path, doc_name)

    print(f"[DOCX] Processed {doc_name}")
    append_log(f"Ingested DOCX: {doc_name}")


def process_documents() -> None:
    docs_dir = Path(DOCUMENTS_DIR)

    if not docs_dir.exists():
        print(f"[WARN] Documents directory not found: {DOCUMENTS_DIR}")
        return

    files_found = False

    for filename in os.listdir(str(docs_dir)):
        path = str(docs_dir / filename)

        if filename.lower().endswith(".pdf"):
            files_found = True
            try:
                ingest_pdf(path)
            except Exception as e:
                print(f"[ERROR] Failed to process PDF {filename}: {e}")
                append_log(f"ERROR: Failed to process PDF {filename}: {e}")

        elif filename.lower().endswith(".docx"):
            files_found = True
            try:
                ingest_docx(path)
            except Exception as e:
                print(f"[ERROR] Failed to process DOCX {filename}: {e}")
                append_log(f"ERROR: Failed to process DOCX {filename}: {e}")

        else:
            print(f"[SKIP] Unknown file type: {filename}")

    if not files_found:
        print(f"[INFO] No documents found in {DOCUMENTS_DIR}")


if __name__ == "__main__":
    process_documents()
